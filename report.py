"""Формирование отчёта в DOCX по всему сеансу работы."""
import hashlib
import io
import os
from datetime import datetime

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

import analysis


def file_sha256(path):
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def build_spectrogram(case):
    """Спектрограмма исследуемой фонограммы."""
    figure, axes = plt.subplots(figsize=(9, 3.0))
    power, freqs, times = analysis.spectrogram_db(case["signal"], case["sr"])
    axes.pcolormesh(times, freqs, power, cmap="magma", shading="auto")
    axes.set_xlabel("Время, с")
    axes.set_ylabel("Частота, Гц")
    figure.tight_layout()

    buffer = io.BytesIO()
    figure.savefig(buffer, dpi=120, format="png")
    plt.close(figure)
    buffer.seek(0)
    return buffer


def build_one_plot(case, name, result):
    """Кривая одного детектора с выделенными участками."""
    figure, axes = plt.subplots(figsize=(9, 2.6))
    axes.plot(result["times"], result["probs"])
    axes.axhline(case["threshold"], linestyle="--", color="red")
    for t0, t1 in result["intervals"]:
        axes.axvspan(t0, t1, alpha=0.3, color="red")
    axes.set_ylim(0, 1)
    axes.set_title(name, fontsize=9)
    axes.set_xlabel("Время, с")
    axes.set_ylabel("Оценка")
    figure.tight_layout()

    buffer = io.BytesIO()
    figure.savefig(buffer, dpi=120, format="png")
    plt.close(figure)
    buffer.seek(0)
    return buffer


def build_plot(case):
    """Все кривые сеанса на одном графике."""
    figure, axes = plt.subplots(figsize=(9, 3.5))
    drawn = False
    for name, result in case["results"].items():
        if len(result["times"]):
            axes.plot(result["times"], result["probs"], label=name)
            drawn = True
    if not drawn:
        plt.close(figure)
        return None

    axes.axhline(case["threshold"], linestyle="--", color="red")
    axes.set_ylim(0, 1)
    axes.set_xlabel("Время, с")
    axes.set_ylabel("Оценка")
    axes.legend(fontsize=7)
    figure.tight_layout()

    buffer = io.BytesIO()
    figure.savefig(buffer, dpi=120, format="png")
    plt.close(figure)
    buffer.seek(0)
    return buffer


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, title in zip(table.rows[0].cells, headers):
        cell.text = title
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = str(value)
    return table


def build_report(out_path, case):
    """case содержит сведения о файле, результаты всех детекторов и журнал."""
    info = case.get("info") or {}
    checks = case["checks"]
    doc = Document()

    doc.add_heading("Заключение по исследованию фонограммы", level=0)
    doc.add_paragraph(f"Дата формирования: {datetime.now():%d.%m.%Y %H:%M}")

    # Объект исследования — исходный файл; для видео это сам видеофайл,
    # а анализировалась извлечённая из него дорожка.
    source = case.get("source") or case["path"]

    doc.add_heading("1. Объект исследования", level=1)
    add_table(doc, ["Параметр", "Значение"], [
        ("Имя файла", os.path.basename(source)),
        ("Размер, байт", os.path.getsize(source)),
        ("SHA-256", file_sha256(source)),
        ("Контейнер / кодек",
         f"{info.get('container', '?')} / {info.get('codec', '?')}"),
        ("Частота дискретизации, Гц", case["sr"]),
        ("Длительность, с", f"{case['duration']:.2f}"),
    ])
    if case.get("source") and case["source"] != case["path"]:
        doc.add_paragraph(
            "Звуковая дорожка извлечена из видеофайла копированием потока, "
            "без перекодирования. Хеш и размер приведены для исходного файла."
        )

    doc.add_heading("2. Применённые детекторы", level=1)
    if case["results"]:
        rows = []
        for name, result in case["results"].items():
            probs = result["probs"]
            verdict = (result.get("service") or {}).get("verdict", "—")
            rows.append((
                name,
                len(probs),
                f"{max(probs):.3f}" if len(probs) else "—",
                len(result["intervals"]),
                verdict,
            ))
        add_table(doc, ["Модель", "Оценок", "Максимум", "Участков", "Вердикт"], rows)
        doc.add_paragraph(f"Порог принятия решения: {case['threshold']}")
    else:
        doc.add_paragraph("Детекторы не запускались.")

    doc.add_heading("3. Оценки во времени", level=1)

    if case.get("signal") is not None:
        doc.add_paragraph("Спектрограмма фонограммы:")
        doc.add_picture(build_spectrogram(case), width=Inches(6))

    plot = build_plot(case)
    if plot is not None and len(case["results"]) > 1:
        doc.add_paragraph("Сводный график по всем моделям:")
        doc.add_picture(plot, width=Inches(6))

    # Отдельный график на каждую запущенную модель
    for name, result in case["results"].items():
        if not len(result["times"]):
            continue
        doc.add_heading(name, level=2)
        doc.add_picture(build_one_plot(case, name, result), width=Inches(6))
        if result["intervals"]:
            doc.add_paragraph("Участки выше порога:")
            for t0, t1 in result["intervals"]:
                doc.add_paragraph(f"{t0:.2f} — {t1:.2f} с", style="List Bullet")
        else:
            doc.add_paragraph("Участков выше порога не выявлено.")

    doc.add_heading("4. Классические признаки обработки", level=1)
    doc.add_paragraph(f"Смещение постоянной составляющей: {checks['dc_offset']:.6f}")
    doc.add_paragraph(f"Частота среза спектра: {checks['cutoff_hz']:.0f} Гц")
    doc.add_paragraph(
        f"Участков с постоянным значением отсчётов: {len(checks['constant_runs'])}")
    doc.add_paragraph(f"Пар повторяющихся фрагментов: {len(checks['repeats'])}")
    if info.get("lossy"):
        doc.add_paragraph(
            "Фонограмма представлена в формате со сжатием с потерями. Частота среза "
            "спектра в этом случае объясняется работой кодека и не свидетельствует "
            "о редактировании. Исследование выполнено по декодированному сигналу."
        )

    doc.add_heading("5. Журнал работы", level=1)
    if case["journal"]:
        add_table(doc, ["Время", "Событие"], case["journal"])
    else:
        doc.add_paragraph("Журнал пуст.")

    doc.add_heading("6. Ограничения", level=1)
    doc.add_paragraph(
        "Результат носит вероятностный характер и не является выводом о подлинности "
        "записи. Прототип не прошёл валидацию по методике экспертного учреждения; "
        "оценки на фонограммах, полученных неизвестными системами синтеза, "
        "могут быть занижены. Модели, возвращающие одну оценку на весь файл, "
        "показаны ровной линией — привязки к времени они не дают."
    )

    doc.save(out_path)
    return out_path
